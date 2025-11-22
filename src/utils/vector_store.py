from __future__ import annotations

import logging
import shutil
from functools import lru_cache
from pathlib import Path
from typing import Iterable, List, Sequence

from docx import Document as DocxDocument
from langchain.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from pypdf import PdfReader

from ..config.config import get_settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def _normalize_dir(path: Path | str | None) -> Path:
    if path is None:
        return Path()
    return Path(path).expanduser().resolve()


def _vector_store_exists(directory: Path) -> bool:
    return directory.exists() and any(directory.iterdir())


def _reset_directory(directory: Path) -> None:
    if directory.exists():
        shutil.rmtree(directory)
    directory.mkdir(parents=True, exist_ok=True)


def _load_pdf(path: Path) -> List[Document]:
    docs: List[Document] = []
    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # pragma: no cover - upstream library edge cases
        logger.warning("Failed to read PDF %s: %s", path, exc)
        return docs

    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(
                Document(
                    page_content=text.strip(),
                    metadata={
                        "source": path.name,
                        "path": str(path),
                        "page": idx,
                        "type": "pdf",
                    },
                )
            )
    return docs


def _load_docx(path: Path) -> List[Document]:
    docs: List[Document] = []
    try:
        doc = DocxDocument(str(path))
    except Exception as exc:  # pragma: no cover
        logger.warning("Failed to read DOCX %s: %s", path, exc)
        return docs

    buffer: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            buffer.append(text)

    joined = "\n".join(buffer).strip()
    if joined:
        docs.append(
            Document(
                page_content=joined,
                metadata={"source": path.name, "path": str(path), "type": "docx"},
            )
        )
    return docs


def _collect_documents(source_dir: Path, extensions: Iterable[str] | None = None) -> List[Document]:
    if extensions is None:
        extensions = SUPPORTED_EXTENSIONS
    allowed = {ext.lower() for ext in extensions}

    documents: List[Document] = []
    if not source_dir.exists():
        logger.warning("Reference docs directory %s does not exist.", source_dir)
        return documents

    for path in sorted(source_dir.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix.lower() not in allowed:
            continue
        if path.suffix.lower() == ".pdf":
            documents.extend(_load_pdf(path))
        elif path.suffix.lower() == ".docx":
            documents.extend(_load_docx(path))
    return documents


@lru_cache
def _get_embeddings() -> HuggingFaceEmbeddings:
    settings = get_settings()
    model_kwargs = {"device": settings.embedding_device}
    encode_kwargs = {"normalize_embeddings": True}
    return HuggingFaceEmbeddings(
        model_name=settings.embedding_model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs,
    )


def build_reference_vector_store(
    docs_dir: Path | str | None = None,
    persist_dir: Path | str | None = None,
    collection_name: str | None = None,
    force_rebuild: bool = False,
) -> Path:
    settings = get_settings()
    source_dir = _normalize_dir(docs_dir) if docs_dir else settings.reference_docs_dir
    storage_dir = _normalize_dir(persist_dir) if persist_dir else settings.vector_store_dir
    collection = collection_name or settings.vector_collection_name

    documents = _collect_documents(source_dir)
    if not documents:
        raise FileNotFoundError(
            f"No reference documents found under {source_dir}. "
            "Add PDF/DOCX files or point --docs-dir to the correct folder."
        )

    if force_rebuild:
        _reset_directory(storage_dir)
    else:
        storage_dir.mkdir(parents=True, exist_ok=True)
        if _vector_store_exists(storage_dir):
            logger.info(
                "Chroma store already exists at %s; skipping rebuild (use --force to overwrite).",
                storage_dir,
            )
            return storage_dir

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.reference_chunk_size,
        chunk_overlap=settings.reference_chunk_overlap,
    )
    split_docs = splitter.split_documents(documents)
    if not split_docs:
        raise RuntimeError("Document splitting produced no content; check the source files.")

    embeddings = _get_embeddings()
    Chroma.from_documents(
        documents=split_docs,
        embedding=embeddings,
        collection_name=collection,
        persist_directory=str(storage_dir),
    )
    logger.info(
        "Persisted %d chunks from %d documents into %s (collection=%s).",
        len(split_docs),
        len(documents),
        storage_dir,
        collection,
    )
    return storage_dir


@lru_cache
def _get_vector_store() -> Chroma | None:
    settings = get_settings()
    persist_dir = settings.vector_store_dir
    if not _vector_store_exists(persist_dir):
        logger.warning(
            "Chroma vector store not found at %s. Run `python -m scripts.build_reference_store` first.",
            persist_dir,
        )
        return None
    embeddings = _get_embeddings()
    return Chroma(
        embedding_function=embeddings,
        collection_name=settings.vector_collection_name,
        persist_directory=str(persist_dir),
    )


def retrieve_reference_chunks(question: str, k: int = 4) -> List[dict]:
    if not question.strip():
        return []
    store = _get_vector_store()
    if store is None:
        return []

    docs = store.similarity_search(question, k=k)
    cleaned: List[dict] = []
    for doc in docs:
        cleaned.append(
            {
                "content": doc.page_content.strip(),
                "source": doc.metadata.get("source"),
                "page": doc.metadata.get("page"),
                "path": doc.metadata.get("path"),
            }
        )
    return cleaned


__all__ = [
    "build_reference_vector_store",
    "retrieve_reference_chunks",
]

